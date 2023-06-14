import pandas as pd
import openai
import os
from docx import Document
import datetime as dt
from dotenv import load_dotenv
import concurrent.futures
# import streamlit as st 

try:
    openai.api_key = st.secrets["OPENAI_API_KEY"]
except:
    load_dotenv()
    openai.api_key = os.getenv("OPENAI_API_KEY")


def create_docx(dicc, school_name, teacher_name, subject):
    """
    Creates a Word document with the given information.

    Args:
        dicc (dict): A dictionary with the topics and questions for the subject.
        school_name (str): The name of the school.
        teacher_name (str): The name of the teacher.
        subject (str): The subject that is being generated for (e.g. Math, Biology).

    Returns:
        None
    """
    
    fecha = dt.datetime.now()
    # Create a new Word document
    document = Document()

    # Set the document properties
    document.core_properties.created = fecha
    document.sections[0].header.paragraphs[0].text = f'{school_name} \t\t{fecha.day}-{fecha.month}-{fecha.year} \Profesor: {teacher_name}'

    # Add the subtitle
    topics = ", ".join(dicc.keys())
    document.add_paragraph(f'Guía de {subject}', style='Title')
    document.add_paragraph(f'Temas revisados: {topics}', style='Subtitle')

    # Add content for each topic and question
    for topic, questions in dicc.items():
        document.add_heading(topic, level=2).bold = True
        for question, content in questions.items():
            document.add_paragraph().add_run(question).bold = True
            document.add_paragraph(content)

    # Save the document
    return document

def generate_guide(excel_file, course, subject, teacher_name, school_name):
    """
    Generates a guide with questions for students based on an Excel file.

    Args:
        excel_file (str): The path to the Excel file containing the subjects and number of questions.
        course (str): The course of the students (e.g. primaria, secundaria).
        subject (int): The subject that is being generated for (e.g. Math, Biology).
        teacher_name (str): The name of the teacher.
        school_name (str): The name of the school.

    Returns:
        None
    """
    # Read Excel file
    df = pd.read_excel(excel_file)
    df.columns = ['topic', 'number']

    df.loc[:,'topic_english'] = translate_topics(df.topic.tolist())
    
    # Initialize a list to store all the questions
    all_questions = {}
    
    # Iterate over each row in the dataframe
    for _, row in df.iterrows():
        topic = row['topic']
        topic_english = row['topic_english']
        num_questions = row['number']
        
        # Generate questions using OpenAI API
        questions = generate_openai_questions(topic_english, num_questions, course, subject)
        
        # Append questions to the list
        all_questions[topic] = questions
   
    # Create a Word document
    doc = create_docx(all_questions, school_name, teacher_name, subject)
    
    # Save the Word document
    doc.save('guide.docx')

def get_course(course):
    """
    Get the course name based on the course number.

    Args:
        course (int): Course number.

    Returns:
        str: Course name.

    Example:
        course_number = 2
        course_name = get_course(course_number)
    """
    if course == 0:
        return 
    elif course == 1:
        return f'{course}st grade'
    elif course == 2:
        return f'{course}nd grade'
    elif course == 3:
        return f'{course}rd grade'
    else:
        return(f'{course}th grade')
    

def generate_openai_questions(topic, num_questions, course, subject):
    """
    Generates questions using the OpenAI API for a given subject.

    Args:
        topic (str): The topic of the questions. (e.g. equations, cells, etc)
        num_questions (int): The number of questions to generate.
        course (str): The course of the students (e.g. primaria, secundaria).
        subject (str): The subject that is being generated for (e.g. Math, Biology).

    Returns:
        questions: A dictionary of generated questions.
    """
    # Get configuration of messages
    messages = generate_messages(topic, course, subject)

    # Create a ThreadPoolExecutor
    with concurrent.futures.ThreadPoolExecutor() as executor:
        # Submit the tasks for generating comments
        results = [executor.submit(generate_question, messages) for _ in range(num_questions)]

        # Retrieve the generated comments
        questions = [result.result() for result in concurrent.futures.as_completed(results)]
       
    # Transform comments into dictionary
    questions = {f'Pregunta {_+1}' : questions[_] for _ in range(len(questions))}

    return questions

def generate_question(messages):
    """
    Generate a question based on a list of conversation messages.

    Args:
        messages (list): List of messages in the conversation.

    Returns:
        str: Generated question.

    Example:
        messages = [
            {"role": "system", "content": "System message"},
            {"role": "user", "content": "User message"}
        ]
        question = generate_question(messages)
    """
    # Call the OpenAI API to generate the response based on the conversation history
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages
    )

    # Extract the generated question from the response and append it to the list
    generated_question = response.choices[0].message.content
    return generated_question

def generate_messages(topic, course, subject):
    """
    Generate a list of messages for a professor designing study guides.

    Args:
        topic (str): The topic for which the multiple-choice question is generated.
        course (str): The course of the students.
        subject (str): The subject of the professor.

    Returns:
        list: List of messages for the conversation.

    Example:
        topic = "algebra"
        course = "10"
        subject = "mathematics"
    """
    #Get the course of the students
    course = get_course(course)

    # Define the initial system message providing context
    context = {
        "role": "system",
        "content": f"You are a great {subject} professor designing study guides for your {course} students. For the guide, you need to generate multiple-choice questions and include explanations on how to solve each question. Write in spanish"
    }
    messages = [context]

    # Prepare the user message as an order prompt for generating a question
    order_prompt = {
        'role': 'user',
        'content': f"Generate a multiple-choice question about {topic}."
    }
    messages.append(order_prompt)
    return messages

def translate_topic(topic):
    # Define the initial system message providing context
    context = {
        "role": "system",
        "content": f"You are a great English-Spanish translator."
    }
    messages = [context]

    # Prepare the user message as an order prompt for translating the topic
    order_prompt = {
        'role': 'user',
        'content': f"Translate the following topic to English: {topic}."
    }
    messages.append(order_prompt)

    # Call the OpenAI API to generate the response based on the conversation history
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages
    )

    # Extract the generated question from the response and append it to the list
    topic = response.choices[0].message.content
    return topic

def translate_topics(list_topics):
    # Create a ThreadPoolExecutor
    with concurrent.futures.ThreadPoolExecutor() as executor:
        # Use map() to submit tasks for generating translation and retrieve the results
        topics = executor.map(translate_topic, list_topics)

    # Return the translated topics
    return list(topics)